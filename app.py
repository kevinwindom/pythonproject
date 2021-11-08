from docx import Document
from docx.shared import Inches

document = Document()

# profile picture
document.add_picture("unnamed.jpeg", width=Inches(2.0), height=Inches(2.0))

# name phone and email
name = input("What is your name: ")
phone_number = input("What is your phone number: ")
email = input("What is your email: ")

document.add_paragraph(
    name + " | " + phone_number + " | " + email)

# about me
document.add_heading("About me")
about_me = input("Tell me about yourself: ")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter previous employer: ")
start_date = input("Start Date: ")
end_date = input("End Date: ")

p.add_run(company + " ").bold = True
p.add_run(start_date + "-" + end_date + "\n").italic = True

experience_details = input(
    "Describe your experience at " + company + ": "
)
p.add_run(experience_details)

# more experience
while True:
    has_more_experiences = input(
        "Do you have more experiences? Yes or No: ")
    if has_more_experiences.upper() == "YES":
        p = document.add_paragraph()

        company = input("Enter previous employer: ")
        start_date = input("Start Date: ")
        end_date = input("End Date: ")

        p.add_run(company + " ").bold = True
        p.add_run(start_date + "-" + end_date + "\n").italic = True

        experience_details = input(
            "Describe your experience at " + company + ": "
        )
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading("Skills")
skill = input("Enter skill: ")
p = document.add_paragraph(skill)
p.style = "List Bullet"

while True:
    more_skills = input("Do you have more skills? Yes or No: ")
    if more_skills.upper() == "YES":
        skill = input("Enter skill: ")
        p = document.add_paragraph(skill)
        p.style = "List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated using PyCharm for this project"


document.save("cv.docx")